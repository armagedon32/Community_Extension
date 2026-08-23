import json
import os

from flask import Blueprint, current_app, flash, redirect, render_template, request, url_for
from flask_login import current_user, login_required
from werkzeug.utils import secure_filename

from app import db
from app.models import DOCUMENT_CATEGORIES, PROJECT_CATEGORIES, Document, MLModel

ml_bp = Blueprint("ml", __name__)

_ALLOWED_EXT = {"txt", "pdf", "doc", "docx"}


def _doc_folder():
    folder = os.path.join(current_app.config["UPLOAD_FOLDER"], "documents")
    return folder


def _save_upload(file):
    os.makedirs(_doc_folder(), exist_ok=True)
    filename = secure_filename(file.filename) or "document.txt"
    file.save(os.path.join(_doc_folder(), filename))
    return filename


def _read_uploaded_text(filename):
    """Extract readable text from an uploaded document.

    Supports plain text (.txt), PDF (.pdf) and Word (.docx). ``.doc`` files
    have no pure-Python text extractor, so they fall back to a best-effort
    raw read that is usually not useful for classification.
    """
    path = os.path.join(_doc_folder(), os.path.basename(filename))
    if not os.path.exists(path):
        return ""

    ext = (os.path.splitext(filename)[1] or "").lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".docx":
            return _extract_docx(path)
        if ext == ".doc":
            with open(path, "rb") as f:
                return f.read().decode("utf-8", errors="ignore")
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def _extract_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(path):
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _delete_doc_file(filename):
    path = os.path.join(_doc_folder(), os.path.basename(filename))
    if os.path.exists(path):
        os.remove(path)


@ml_bp.route("/ml/documents")
@login_required
def list_documents():
    _scan_unclassified()

    from sqlalchemy import or_

    type_filter = (request.args.get("type") or "").strip()
    domain_filter = (request.args.get("domain") or "").strip()

    query = Document.query
    if type_filter:
        query = query.filter(
            or_(Document.category == type_filter, Document.predicted_category == type_filter)
        )
    if domain_filter:
        query = query.filter(
            or_(Document.domain == domain_filter, Document.predicted_domain == domain_filter)
        )

    documents = query.order_by(Document.created_at.desc()).all()

    from app.ml.engine import extract_objective
    objectives = {
        doc.id: extract_objective(doc.content) for doc in documents if doc.content
    }

    return render_template(
        "ml/documents.html",
        documents=documents,
        objectives=objectives,
        type_filter=type_filter,
        domain_filter=domain_filter,
        DOCUMENT_CATEGORIES=DOCUMENT_CATEGORIES,
        PROJECT_CATEGORIES=PROJECT_CATEGORIES,
    )


def _predict_missing(doc):
    """Predict any field that a document is still missing.

    A document may be labeled for its *document type* but have no *project
    category* (or vice versa). This fills the empty one without overwriting
    the existing label.
    """
    from app.ml.engine import classify_text, classify_domain

    if not doc.content:
        return
    _ensure_models()
    if doc.category is None and not doc.predicted_category:
        try:
            pred, _ = classify_text(doc.content)
            if pred:
                doc.predicted_category = pred
        except Exception:
            pass
    if doc.domain is None and not doc.predicted_domain:
        try:
            p_dom, _ = classify_domain(doc.content)
            if p_dom:
                doc.predicted_domain = p_dom
        except Exception:
            pass


def _scan_unclassified():
    """Auto-scan so every document shows a project category on page load.

    Re-reads the content from the stored file for any document whose content
    is empty (e.g. uploaded before PDF text extraction existed), then predicts
    the *project category* (and type) from that scanned content.
    """
    docs = (
        Document.query
        .order_by(Document.created_at.desc())
        .all()
    )
    for doc in docs:
        if not doc.content and doc.filename:
            doc.content = _read_uploaded_text(doc.filename)
        if not (doc.domain or doc.predicted_domain) or not (doc.category or doc.predicted_category):
            _predict_missing(doc)
    db.session.commit()


@ml_bp.route("/ml/documents/new", methods=["GET", "POST"])
@login_required
def add_document():
    if request.method == "POST":
        title = request.form.get("title", "").strip()
        category = request.form.get("category", "").strip()
        domain = request.form.get("domain", "").strip()
        content = request.form.get("content", "").strip()
        is_training = True if request.form.get("is_training") else False
        filename = None

        file = request.files.get("file")
        if file and file.filename:
            filename = _save_upload(file)
            if not content:
                content = _read_uploaded_text(filename)

        if not title or not content:
            flash("Document title and content are required.", "danger")
            return render_template(
                "ml/document_form.html",
                document=None,
                DOCUMENT_CATEGORIES=DOCUMENT_CATEGORIES,
                PROJECT_CATEGORIES=PROJECT_CATEGORIES,
            )

        doc = Document(
            title=title,
            category=category if category else None,
            domain=domain if domain else None,
            content=content,
            filename=filename,
            is_training=is_training,
            uploaded_by=current_user.id,
        )
        db.session.add(doc)
        db.session.commit()

        _predict_missing(doc)
        db.session.commit()

        if is_training and (category or domain):
            flash("Labeled document added. Retrain the model to include it.", "success")
        elif doc.predicted_category or doc.predicted_domain:
            parts = [p for p in (doc.predicted_category, doc.predicted_domain) if p]
            flash(f"Document scanned. Project category detected: {parts[-1]}", "success")
        else:
            flash("Document saved. Could not auto-detect a category for this content yet.", "warning")
        return redirect(url_for("ml.list_documents"))

    return render_template(
        "ml/document_form.html",
        document=None,
        DOCUMENT_CATEGORIES=DOCUMENT_CATEGORIES,
        PROJECT_CATEGORIES=PROJECT_CATEGORIES,
    )


@ml_bp.route("/ml/documents/<int:doc_id>/delete", methods=["POST"])
@login_required
def delete_document(doc_id):
    doc = db.get_or_404(Document, doc_id)
    if doc.filename:
        _delete_doc_file(doc.filename)
    db.session.delete(doc)
    db.session.commit()
    flash("Document deleted.", "info")
    return redirect(url_for("ml.list_documents"))


@ml_bp.route("/ml/train", methods=["POST"])
@login_required
def train_model():
    from app.ml import engine

    type_docs = (
        Document.query
        .filter(Document.is_training.is_(True), Document.category.isnot(None))
        .all()
    )
    if len(type_docs) < 3:
        flash("At least 3 labeled training documents are required.", "danger")
        return redirect(url_for("ml.dashboard"))

    # --- Document type model ---
    texts = [doc.content or "" for doc in type_docs]
    labels = [doc.category for doc in type_docs]
    model, vectorizer, metrics = engine.train_naive_bayes(texts, labels)
    engine.save_model(model, vectorizer, metrics["classes"], "type")

    _record_model_run("Multinomial Naive Bayes - Document Type", metrics)

    # --- Project category/domain model ---
    domain_docs = (
        Document.query
        .filter(Document.is_training.is_(True), Document.domain.isnot(None))
        .all()
    )
    if len(domain_docs) >= 3:
        domain_texts = [doc.content or "" for doc in domain_docs]
        domain_labels = [doc.domain for doc in domain_docs]
        d_model, d_vec, d_metrics = engine.train_domain_model(domain_texts, domain_labels)
        _record_model_run("Multinomial Naive Bayes - Project Category", d_metrics)
        flash(
            f"Models trained. Document Type: Accuracy {metrics['accuracy']:.0%} | "
            f"Project Category: Accuracy {d_metrics['accuracy']:.0%}",
            "success",
        )
        from app.routes.notifications import notify
        notify(
            f"ML models trained — Document Type accuracy {metrics['accuracy']:.0%}, Project Category accuracy {d_metrics['accuracy']:.0%}.",
            category="success",
            link=url_for("ml.dashboard"),
        )
    else:
        flash(
            f"Document type model trained (Accuracy: {metrics['accuracy']:.0%}). "
            f"Add at least 3 documents with a project category/domain to train the second model.",
            "success",
        )
        from app.routes.notifications import notify
        notify(
            f"Document type model trained — accuracy {metrics['accuracy']:.0%}.",
            category="success",
            link=url_for("ml.dashboard"),
        )
    return redirect(url_for("ml.dashboard"))


def _record_model_run(name, metrics):
    model_run = MLModel(
        name=name,
        model_type="Multinomial Naive Bayes",
        status="Trained",
        accuracy=metrics.get("accuracy"),
        precision=metrics.get("precision"),
        recall=metrics.get("recall"),
        f1=metrics.get("f1"),
        samples=metrics.get("samples", 0),
        classes=", ".join(metrics.get("classes") or []),
        metrics_json=json.dumps(metrics, default=str),
    )
    db.session.add(model_run)
    db.session.commit()


@ml_bp.route("/ml/classify/<int:doc_id>", methods=["POST"])
@login_required
def classify_document(doc_id):
    doc = db.get_or_404(Document, doc_id)
    _classify_and_store(doc)
    return redirect(url_for("ml.list_documents"))


@ml_bp.route("/ml/documents/<int:doc_id>/print")
@login_required
def print_document(doc_id):
    from datetime import datetime

    doc = db.get_or_404(Document, doc_id)
    return render_template(
        "ml/document_print.html",
        doc=doc,
        generated_at=datetime.now().strftime("%B %d, %Y %I:%M %p"),
    )


def _ensure_models():
    """Train any missing classifier on-demand so classification always works.

    Returns True if both requested classifiers (type + domain) are ready.
    """
    from app.ml import engine

    trained = {}
    docs = (
        Document.query
        .filter(Document.is_training.is_(True), Document.category.isnot(None))
        .all()
    )
    if engine.load_model("type")[0] is None and len(docs) >= 3:
        model, vectorizer, metrics = engine.train_naive_bayes(
            [d.content or "" for d in docs],
            [d.category for d in docs],
        )
        engine.save_model(model, vectorizer, metrics["classes"], "type")
        trained["type"] = True

    domain_docs = (
        Document.query
        .filter(Document.is_training.is_(True), Document.domain.isnot(None))
        .all()
    )
    if engine.load_model("domain")[0] is None and len(domain_docs) >= 3:
        d_model, d_vec, d_metrics = engine.train_domain_model(
            [d.content or "" for d in domain_docs],
            [d.domain for d in domain_docs],
        )
        engine.save_model(d_model, d_vec, d_metrics["classes"], "domain")
        trained["domain"] = True

    return trained


def _classify_and_store(doc):
    from app.ml.engine import classify_text, classify_domain

    _ensure_models()

    predicted, scores = classify_text(doc.content)
    if predicted:
        doc.predicted_category = predicted
    predicted_domain, _d_scores = classify_domain(doc.content)
    if predicted_domain:
        doc.predicted_domain = predicted_domain
    db.session.commit()

    parts = [p for p in (predicted, predicted_domain) if p]
    if parts:
        flash(f"Document classified as: {' / '.join(parts)}", "success")
        from app.routes.notifications import notify
        notify(
            f"Document '{doc.title}' classified as {' / '.join(parts)}.",
            category="success",
            link=url_for("ml.list_documents"),
        )
    else:
        flash("Unable to classify this document (not enough training data yet).", "warning")


@ml_bp.route("/ml")
@login_required
def dashboard():
    models = MLModel.query.order_by(MLModel.created_at.desc()).all()
    labeled_count = Document.query.filter(Document.is_training.is_(True), Document.category.isnot(None)).count()
    domain_count = Document.query.filter(Document.is_training.is_(True), Document.domain.isnot(None)).count()
    doc_count = Document.query.count()

    def _chart(model):
        if model is None or not model.metrics_json:
            return None
        try:
            m = json.loads(model.metrics_json)
            report = m.get("classification_report") or {}
            return {
                "accuracy": m.get("accuracy"),
                "precision": m.get("precision"),
                "recall": m.get("recall"),
                "f1": m.get("f1"),
                "confusion_matrix": m.get("confusion_matrix"),
                "classes": m.get("classes"),
                "test_samples": m.get("test_samples"),
                "classification_report": report,
            }
        except Exception:
            return None

    type_model = next((m for m in models if "Document Type" in m.name), None)
    domain_model = next((m for m in models if "Project Category" in m.name), None)

    return render_template(
        "ml/index.html",
        models=models,
        latest=type_model,
        latest_domain=domain_model,
        chart=_chart(type_model),
        chart_domain=_chart(domain_model),
        labeled=labeled_count,
        domain_count=domain_count,
        doc_count=doc_count,
        DOCUMENT_CATEGORIES=DOCUMENT_CATEGORIES,
        PROJECT_CATEGORIES=PROJECT_CATEGORIES,
    )